import base64
import json
import sys
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

# Allow running pytest from backend/tests without PYTHONPATH tweaks.
PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from backend.parser import ParserError, normalize_input


def test_normalize_text_chunks_with_line_numbers() -> None:
    content = "line-1\nline-2\nline-3"

    normalized = normalize_input("text", content, chunk_size=2)

    assert normalized["content_type"] == "text"
    assert len(normalized["lines"]) == 3
    assert normalized["lines"][0]["line_number"] == 1
    assert normalized["lines"][2]["line_number"] == 3
    assert len(normalized["chunks"]) == 2
    assert normalized["chunks"][0]["start_line"] == 1
    assert normalized["chunks"][0]["end_line"] == 2
    assert normalized["chunks"][1]["start_line"] == 3


def test_normalize_file_payload_text_base64() -> None:
    raw_text = "alpha\nbeta"
    payload = {
        "file_name": "app.log",
        "content_base64": base64.b64encode(raw_text.encode("utf-8")).decode("utf-8"),
    }

    normalized = normalize_input("file", json.dumps(payload), chunk_size=10)

    assert normalized["source"] == "file"
    assert normalized["content_type"] == "logs"
    assert normalized["text"] == raw_text
    assert normalized["metadata"]["file_suffix"] == ".log"


def test_invalid_base64_raises_parser_error() -> None:
    payload = {
        "file_name": "bad.txt",
        "content_base64": "not-base64",
    }

    with pytest.raises(ParserError) as exc:
        normalize_input("file", json.dumps(payload))

    assert exc.value.code == "invalid_base64"


def test_pdf_payload_extracts_without_crash() -> None:
    pdf = PdfWriter()
    pdf.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    pdf.write(buffer)

    payload = {
        "file_name": "report.pdf",
        "content_base64": base64.b64encode(buffer.getvalue()).decode("utf-8"),
    }

    normalized = normalize_input("file", json.dumps(payload), chunk_size=10)

    assert normalized["content_type"] == "file"
    assert normalized["metadata"]["file_suffix"] == ".pdf"
    assert normalized["metadata"]["extraction_method"] == "pypdf"


def test_docx_payload_extracts_text() -> None:
    doc = Document()
    doc.add_paragraph("username=admin")
    doc.add_paragraph("password=secret123")

    buffer = BytesIO()
    doc.save(buffer)

    payload = {
        "file_name": "credentials.docx",
        "content_base64": base64.b64encode(buffer.getvalue()).decode("utf-8"),
    }

    normalized = normalize_input("file", json.dumps(payload), chunk_size=10)

    assert normalized["metadata"]["file_suffix"] == ".docx"
    assert normalized["metadata"]["extraction_method"] == "python-docx"
    assert "password=secret123" in normalized["text"]


def test_unsupported_file_type_raises_parser_error() -> None:
    payload = {
        "file_name": "archive.zip",
        "content_base64": base64.b64encode(b"dummy").decode("utf-8"),
    }

    with pytest.raises(ParserError) as exc:
        normalize_input("file", json.dumps(payload))

    assert exc.value.code == "unsupported_file_type"


def test_invalid_file_payload_json_raises_parser_error() -> None:
    with pytest.raises(ParserError) as exc:
        normalize_input("file", "not-json")

    assert exc.value.code == "invalid_file_payload"


def test_large_file_chunking_is_stable() -> None:
    content = "\n".join([f"line-{index}" for index in range(1, 10001)])

    normalized = normalize_input("log", content, chunk_size=50)

    assert normalized["metadata"]["line_count"] == 10000
    assert normalized["metadata"]["chunk_count"] == 200
    assert normalized["chunks"][0]["start_line"] == 1
    assert normalized["chunks"][0]["end_line"] == 50
    assert normalized["chunks"][-1]["start_line"] == 9951
    assert normalized["chunks"][-1]["end_line"] == 10000
